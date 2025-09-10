import os
from flask import Flask, request, render_template, flash, redirect, url_for, jsonify
from flask_wtf import FlaskForm, CSRFProtect
from flask_wtf.file import FileField, FileRequired, FileAllowed
from wtforms import SubmitField
from werkzeug.utils import secure_filename
import docx
import PyPDF2
from io import BytesIO
import logging
from datetime import datetime
import hashlib
import mimetypes
from config import config

# Initialize Flask app with configuration
def create_app(config_name='default'):
    app = Flask(__name__)
    app.config.from_object(config[config_name])
    
    # Initialize CSRF protection
    csrf = CSRFProtect(app)
    
    # Configure logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    logger = logging.getLogger(__name__)
    
    return app, logger

app, logger = create_app(os.environ.get('FLASK_ENV', 'development'))

# File upload form with CSRF protection
class FileUploadForm(FlaskForm):
    file = FileField('Document File', validators=[
        FileRequired(message='Please select a file'),
        FileAllowed(['docx', 'pdf'], 'Only DOCX and PDF files are allowed')
    ])
    submit = SubmitField('Extract Metadata')

def validate_file_content(file_stream, filename):
    """Validate file content matches extension"""
    file_stream.seek(0)
    mime_type, _ = mimetypes.guess_type(filename)
    
    # Read first few bytes to check file signature
    header = file_stream.read(8)
    file_stream.seek(0)
    
    if filename.lower().endswith('.pdf'):
        if not header.startswith(b'%PDF'):
            raise ValueError("File content doesn't match PDF format")
    elif filename.lower().endswith('.docx'):
        if not header.startswith(b'PK'):  # ZIP signature for DOCX
            raise ValueError("File content doesn't match DOCX format")
    
    return True

def calculate_file_hash(file_stream):
    """Calculate SHA-256 hash of file for integrity"""
    file_stream.seek(0)
    hash_sha256 = hashlib.sha256()
    for chunk in iter(lambda: file_stream.read(4096), b""):
        hash_sha256.update(chunk)
    file_stream.seek(0)
    return hash_sha256.hexdigest()

def format_date(date_obj):
    """Format date object to readable string"""
    if date_obj:
        if isinstance(date_obj, datetime):
            return date_obj.strftime('%Y-%m-%d %H:%M:%S')
        return str(date_obj)
    return 'Not available'

def format_file_size(size_bytes):
    """Format file size in human readable format"""
    if size_bytes == 0:
        return "0 B"
    size_names = ["B", "KB", "MB", "GB"]
    i = 0
    while size_bytes >= 1024 and i < len(size_names) - 1:
        size_bytes /= 1024.0
        i += 1
    return f"{size_bytes:.1f} {size_names[i]}"

def extract_metadata_from_docx(doc, file_stream):
    """Extract comprehensive metadata from DOCX file"""
    metadata = {}
    core_props = doc.core_properties
    
    # Basic metadata
    metadata['Title'] = getattr(core_props, 'title', None) or 'Not specified'
    metadata['Subject'] = getattr(core_props, 'subject', None) or 'Not specified'
    metadata['Author'] = getattr(core_props, 'creator', None) or 'Not specified'
    metadata['Keywords'] = getattr(core_props, 'keywords', None) or 'Not specified'
    metadata['Last Modified By'] = getattr(core_props, 'last_modified_by', None) or 'Not specified'
    metadata['Revision'] = getattr(core_props, 'revision', None) or 'Not specified'
    metadata['Modified Date'] = format_date(getattr(core_props, 'modified', None))
    metadata['Created Date'] = format_date(getattr(core_props, 'created', None))
    metadata['Category'] = getattr(core_props, 'category', None) or 'Not specified'
    metadata['Comments'] = getattr(core_props, 'comments', None) or 'Not specified'
    
    # Document structure
    metadata['Number of Paragraphs'] = len(doc.paragraphs)
    metadata['Number of Tables'] = len(doc.tables)
    
    # Count sections
    sections = doc.sections
    metadata['Number of Sections'] = len(sections)
    
    # Word and character count
    word_count = 0
    char_count = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text
        word_count += len(text.split())
        char_count += len(text)
    
    metadata['Approximate Word Count'] = word_count
    metadata['Character Count'] = char_count
    
    # File integrity
    metadata['File Hash (SHA-256)'] = calculate_file_hash(file_stream)
    
    return metadata

def extract_metadata_from_pdf(file_stream):
    """Extract comprehensive metadata from PDF file"""
    metadata = {}
    try:
        pdf_reader = PyPDF2.PdfReader(file_stream)
        info = pdf_reader.metadata
        
        if info:
            metadata['Title'] = getattr(info, 'title', None) or 'Not specified'
            metadata['Author'] = getattr(info, 'author', None) or 'Not specified'
            metadata['Subject'] = getattr(info, 'subject', None) or 'Not specified'
            metadata['Creator'] = getattr(info, 'creator', None) or 'Not specified'
            metadata['Producer'] = getattr(info, 'producer', None) or 'Not specified'
            metadata['Creation Date'] = format_date(getattr(info, 'creation_date', None))
            metadata['Modification Date'] = format_date(getattr(info, 'modification_date', None))
        else:
            metadata['Title'] = 'Not available'
            metadata['Author'] = 'Not available'
            metadata['Subject'] = 'Not available'
            metadata['Creator'] = 'Not available'
            metadata['Producer'] = 'Not available'
            metadata['Creation Date'] = 'Not available'
            metadata['Modification Date'] = 'Not available'
        
        metadata['Number of Pages'] = len(pdf_reader.pages)
        
        # Check if PDF is encrypted
        metadata['Encrypted'] = 'Yes' if pdf_reader.is_encrypted else 'No'
        
        # PDF version
        if hasattr(pdf_reader, 'pdf_header'):
            metadata['PDF Version'] = pdf_reader.pdf_header
        
        # File integrity
        metadata['File Hash (SHA-256)'] = calculate_file_hash(file_stream)
        
    except Exception as e:
        logger.error(f"Error extracting PDF metadata: {str(e)}")
        raise
    
    return metadata

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    form = FileUploadForm()
    
    if form.validate_on_submit():
        try:
            file = form.file.data
            filename = secure_filename(file.filename)
            
            # Get file size
            file.seek(0, 2)  # Seek to end
            file_size = file.tell()
            file.seek(0)  # Reset to beginning
            
            file_bytes = BytesIO(file.read())
            
            # Validate file content
            validate_file_content(file_bytes, filename)
            
            if filename.lower().endswith('.docx'):
                doc = docx.Document(file_bytes)
                metadata = extract_metadata_from_docx(doc, file_bytes)
                file_type = 'Microsoft Word Document (DOCX)'
            else:  # PDF
                file_bytes.seek(0)
                metadata = extract_metadata_from_pdf(file_bytes)
                file_type = 'Portable Document Format (PDF)'
            
            # Add file information
            metadata['File Name'] = filename
            metadata['File Type'] = file_type
            metadata['File Size'] = format_file_size(file_size)
            
            return render_template('metadata.html', 
                                 metadata=metadata, 
                                 filename=filename,
                                 file_size=format_file_size(file_size))
            
        except ValueError as e:
            flash(f'File validation error: {str(e)}', 'error')
            logger.warning(f"File validation failed for {filename}: {str(e)}")
        except Exception as e:
            logger.error(f"Error processing file {filename}: {str(e)}")
            flash(f'Error processing the file: {str(e)}', 'error')
    
    return render_template('upload.html', form=form)

@app.route('/api/upload-progress')
def upload_progress():
    """API endpoint for upload progress (placeholder for future implementation)"""
    return jsonify({'progress': 100, 'status': 'complete'})

@app.errorhandler(413)
def too_large(e):
    flash('File is too large. Maximum size is 16MB.', 'error')
    return redirect(url_for('upload_file'))

@app.errorhandler(400)
def bad_request(e):
    flash('Bad request. Please check your file and try again.', 'error')
    return redirect(url_for('upload_file'))

@app.errorhandler(500)
def internal_error(e):
    logger.error(f"Internal server error: {str(e)}")
    flash('An internal error occurred. Please try again later.', 'error')
    return redirect(url_for('upload_file'))

if __name__ == "__main__":
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('FLASK_ENV') == 'development'
    app.run(debug=debug, host='0.0.0.0', port=port)